import streamlit as st
import numpy as np
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
import os
from dotenv import load_dotenv
from rags import RAG_IMPLEMENTATIONS, RAG_DESCRIPTIONS, process_query
import fitz
from docx import Document as DocxDocument
import markdown
from bs4 import BeautifulSoup
import mimetypes
import whisper
import requests


load_dotenv()

# API Key Handling (use secrets if possible)
GOOGLE_API_KEY = os.getenv('GEMINI_API_KEY')
videoTotext = whisper.load_model("base")

if GOOGLE_API_KEY is None:
    GOOGLE_API_KEY = st.text_input("Enter your Gemini API key:", type="password")

os.environ['GOOGLE_API_KEY'] = GOOGLE_API_KEY

# Initialize OpenAI client and LLM
llm = ChatGoogleGenerativeAI(model="gemini-1.5-pro")
embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

# Initialize session state variables
if 'paragraphs' not in st.session_state:
    st.session_state.paragraphs = []
if 'paragraph_embeddings' not in st.session_state:
    st.session_state.paragraph_embeddings = np.array([])
if 'query_embedding' not in st.session_state:
    st.session_state.query_embedding = None
if 'text_processed' not in st.session_state:
    st.session_state.text_processed = False

# Streamlit App
st.title("Retrieval Augmented Generation (RAG) Demo")

# Add RAG type selector
rag_type = st.selectbox(
    "Select RAG Implementation:",
    list(RAG_IMPLEMENTATIONS.keys()),
    help="Choose the type of RAG implementation to use"
)

url = st.text_input("Enter a webpage URL to process:")
if url:
    st.write(f"URL Provided: {url}")

# Functions to process different file types
def read_pdf(file):
    pdf_document = fitz.open(stream=file.read(), filetype="pdf")
    text = ""
    for page_num in range(len(pdf_document)):
        page = pdf_document[page_num]
        text += page.get_text()
    return text


def read_docx(file):
    doc = DocxDocument(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text


def read_markdown(file):
    md_text = file.read().decode("utf-8")
    html = markdown.markdown(md_text)
    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text()

def read_html(file):
    content = file.read().decode("utf-8", errors="ignore")
    soup = BeautifulSoup(content, "html.parser")
    text = soup.get_text(separator="\n")
    return text

# def read_video(file):
#     temp_video_path = "temp_video.mp4"
#     with open(temp_video_path, "wb") as f:
#         f.write(file.read())
#
#     model = whisper.load_model("base")
#     result = model.transcribe(temp_video_path)
#
#     os.remove(temp_video_path)
#
#     return result["text"]
def read_video(file):
    """
    Reads and transcribes a video file using the Whisper model.
    """
    video_folder = "my_study_video"
    os.makedirs(video_folder, exist_ok=True)

    video_path = os.path.join(video_folder, file.name)

    # Save the uploaded video file
    with open(video_path, "wb") as video_file:
        video_file.write(file.read())

    try:
        # Transcribe video using Whisper
        result = videoTotext.transcribe(video_path, word_timestamps=True)

        # Process transcription result
        video_text = "\n\n".join(
            f"[{word_info['start']:.2f}s-{word_info['end']:.2f}s] {word_info['text']}"
            for word_info in result.get('segments', [])
        )

        return video_text

    except Exception as e:
        st.error(f"Error transcribing video: {e}")
        return ""

    finally:
        # Ensure the video file is removed after processing
        try:
            if os.path.exists(video_path):
                os.remove(video_path)
        except Exception as cleanup_error:
            st.warning(f"Could not delete temporary file: {cleanup_error}")


def read_file(uploaded_file):
    file_type, _ = mimetypes.guess_type(uploaded_file.name)

    if file_type == "text/plain":
        return uploaded_file.read().decode("utf-8")
    elif file_type == "application/pdf":
        return read_pdf(uploaded_file)
    # elif file_type == "text/html":
    #     return read_html(uploaded_file)
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return read_docx(uploaded_file)
    elif file_type == "text/markdown":
        return read_markdown(uploaded_file)
    elif file_type and file_type.startswith("video/"):
        return read_video(uploaded_file)
    else:
        raise ValueError("Unsupported file type!")

# Display information about the selected RAG type
st.info(RAG_DESCRIPTIONS[rag_type])

uploaded_files = st.file_uploader("Upload files (text, PDF, Word, Markdown, video)", type=["txt", "pdf", "docx", "md", "mp4", "avi", "mkv"], accept_multiple_files=True)

st.session_state.combined_text = ""
# if uploaded_files:
#     for uploaded_file in uploaded_files:
#         text = uploaded_file.read().decode("utf-8")
#         st.session_state.combined_text += text
#
#     st.text_area("Uploaded Text:", st.session_state.combined_text, height=150)
if uploaded_files:
    for uploaded_file in uploaded_files:
        try:
            text = read_file(uploaded_file)
            st.session_state.combined_text += text
        except ValueError as e:
            st.error(str(e))

    # st.text_area("Uploaded Text:", st.session_state.combined_text, height=150)
if url:
    try:
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        url_text = soup.get_text(separator="\n")
        st.session_state.combined_text += url_text
    except requests.exceptions.RequestException as e:
        st.error(f"Error fetching URL: {e}")

system_message = {
    "role": "system",
    "content": """You are an AI assistant specifically designed to work with RAG (Retrieval Augmented Generation) systems. Your responses must be strictly based on the provided context. Follow these rules:

1. Only answer using information explicitly present in the given context
2. If the context is not relevant to the query, respond with "The provided context is not relevant to answer this query."
3. If the context is partially relevant but incomplete, respond with "The context is incomplete to fully answer this query. Based on the available context, I can only tell you that: [insert relevant information from context]"
4. Do not make assumptions or add information beyond what's in the context
5. If you quote or reference specific parts of the context, indicate this in your response
6. If the context is empty or missing, respond with "No context was provided to answer this query."

Remember: You are a context-dependent system. Never draw from general knowledge - only use the specific context provided in each query."""
}

if rag_type=="Naive RAG":

    chunk_size = st.number_input("Chunk Size:", min_value=1, value=512)
    query = st.text_input("Enter your query:")
    n_results = st.number_input("Number of similar chunks to find:", min_value=1, value=1)
    use_context = st.checkbox("Use Context", value=True)
    # Buttons in one row with three columns
    _, col1, col2, col3, _ = st.columns(5)

    with col1:
        process_button = st.button("Process Text")
    with col2:
        search_button = st.button("Search")
    with col3:
        run_query_button = st.button("Run Query")

    # Output section spans full width
    if process_button:
        if not st.session_state.combined_text:
            st.info("Please upload text first.")
        elif not query:
            st.info("Please enter a query first.")
        else:
            process_function = RAG_IMPLEMENTATIONS[rag_type]
            st.session_state.paragraphs, st.session_state.paragraph_embeddings, st.session_state.query_embedding = process_function(
                st.session_state.combined_text, 
                chunk_size,
                embeddings,
                query
            )
            st.session_state.text_processed = True
            st.success(f"Text processed using {rag_type}!")


    if search_button:
        if not st.session_state.text_processed:
            st.info("Please process the text first.")
        elif not query:
            st.info("Please enter a query.")
        else:
            most_similar_indices, _ = process_query(
                query, 
                embeddings,
                st.session_state.paragraph_embeddings,
                st.session_state.query_embedding,
                n_results
            )
            st.subheader("Search Results:")
            for index in most_similar_indices:
                st.write(st.session_state.paragraphs[index])

    if run_query_button:
        if not st.session_state.text_processed:
            st.info("Please process the text first.")
        elif not query:
            st.info("Please enter a query.")
        else:
            most_similar_indices, _ = process_query(
                query, 
                embeddings,
                st.session_state.paragraph_embeddings,
                st.session_state.query_embedding,
                n_results
            )

            if use_context:
                most_similar_paragraph = "\n\n".join([st.session_state.paragraphs[index] for index in most_similar_indices])
                llm_query = f"CONTEXT: {most_similar_paragraph}\n\nQUERY: {query}"
                messages = [
                    system_message,
                    {"role": "user", "content": llm_query},
                ]
            else:
                llm_query = query

                messages = [
                    {
                    "role": "system",
                    "content": "You are a helpful assistant."
                },
                    {"role": "user", "content": llm_query},
                ]

            try:
                response = llm.invoke(messages).content
                st.write(response)

                if use_context:
                    st.subheader("Relevant Context:")
                    for index in most_similar_indices:
                        st.write(st.session_state.paragraphs[index])

            except Exception as e:
                st.error(f"Error generating response: {e}")

else:
    
    query = st.text_input("Enter your query:")
    _, _, col1, _, _ = st.columns(5)
    
    with col1:
        run_query_button = st.button("Run Query")

    # Output section spans full width
    if run_query_button:
        if not st.session_state.combined_text:
            st.info("Please upload text first.")
        elif not query:
            st.info("Please enter a query first.")
        else:
            process_function = RAG_IMPLEMENTATIONS[rag_type]
            response = process_function(
                st.session_state.combined_text, 
                embeddings,
                query,
                llm
            )
            st.write(response)
            st.success(f"Query processed using {rag_type}!")
